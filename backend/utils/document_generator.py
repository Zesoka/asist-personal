import io
from fpdf import FPDF
import docx

def generate_md(markdown_text):
    """Converts markdown text to bytes for file download."""
    return markdown_text.encode("utf-8")

def generate_pdf(markdown_text):
    """Converts markdown text to a basic PDF using fpdf2."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    
    in_code_block = False
    
    for line in markdown_text.split("\n"):
        line = line.rstrip()
        safe_line = line.encode('latin-1', 'replace').decode('latin-1')
        
        if safe_line.startswith("```"):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            pdf.set_font("Courier", size=10)
            pdf.cell(0, 6, safe_line, ln=True)
            continue
            
        if safe_line.startswith("# "):
            pdf.ln(5)
            pdf.set_font("Helvetica", style="B", size=18)
            pdf.cell(0, 10, safe_line[2:], ln=True)
            pdf.ln(2)
        elif safe_line.startswith("## "):
            pdf.ln(4)
            pdf.set_font("Helvetica", style="B", size=15)
            pdf.cell(0, 8, safe_line[3:], ln=True)
            pdf.ln(1)
        elif safe_line.startswith("### "):
            pdf.ln(3)
            pdf.set_font("Helvetica", style="B", size=12)
            pdf.cell(0, 6, safe_line[4:], ln=True)
            pdf.ln(1)
        elif safe_line.startswith("- ") or safe_line.startswith("* "):
            pdf.set_font("Helvetica", size=11)
            pdf.cell(5, 6, "", ln=False)
            pdf.cell(0, 6, chr(149) + " " + safe_line[2:], ln=True)
        elif safe_line.startswith("1. ") or (len(safe_line) > 2 and safe_line[0].isdigit() and safe_line[1:3] == ". "):
            pdf.set_font("Helvetica", size=11)
            pdf.cell(5, 6, "", ln=False)
            pdf.cell(0, 6, safe_line, ln=True)
        else:
            if safe_line.strip() == "":
                pdf.ln(3)
            else:
                pdf.set_font("Helvetica", size=11)
                pdf.multi_cell(0, 6, safe_line)
                
    pdf_bytes = pdf.output()
    if isinstance(pdf_bytes, str):
        return pdf_bytes.encode('latin-1')
    return pdf_bytes

def generate_docx(markdown_text):
    """Converts markdown text to a Word document using python-docx."""
    doc = docx.Document()
    in_code_block = False
    code_text = []
    
    for line in markdown_text.split("\n"):
        line = line.rstrip()
        
        if line.startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = docx.shared.Inches(0.5)
                run = p.add_run("\n".join(code_text))
                run.font.name = 'Courier New'
                run.font.size = docx.shared.Pt(9.5)
                code_text = []
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            code_text.append(line)
            continue
            
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith("1. ") or (len(line) > 2 and line[0].isdigit() and line[1:3] == ". "):
            prefix_len = line.find(" ") + 1
            doc.add_paragraph(line[prefix_len:], style='List Number')
        else:
            if line.strip() != "":
                doc.add_paragraph(line)
                
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()
